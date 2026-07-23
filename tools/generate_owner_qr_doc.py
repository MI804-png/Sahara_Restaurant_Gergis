from pathlib import Path

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "print-assets"
    out_dir.mkdir(parents=True, exist_ok=True)

    # Replace this after Render deploy with your real public URL.
    website_url = "https://YOUR-RENDER-SERVICE.onrender.com"

    qr_path = out_dir / "sahara-restaurant-website-qr.png"
    docx_path = out_dir / "Sahara-Restaurant-Website-QR.docx"

    qr = qrcode.QRCode(version=2, box_size=12, border=2)
    qr.add_data(website_url)
    qr.make(fit=True)
    image = qr.make_image(fill_color="black", back_color="white")
    image.save(qr_path)

    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Sahara Restaurant - Website QR")
    title_run.bold = True
    title_run.font.size = Pt(22)

    document.add_paragraph("")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Scan this QR code to open the website on mobile.").font.size = Pt(12)

    url_line = document.add_paragraph()
    url_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url_line.add_run(website_url).font.size = Pt(10)

    document.add_paragraph("")

    qr_para = document.add_paragraph()
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_para.add_run().add_picture(str(qr_path), width=Inches(3.2))

    document.add_paragraph("")

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Owner copy - ready to print").italic = True

    document.save(docx_path)

    print(f"QR created: {qr_path}")
    print(f"DOCX created: {docx_path}")


if __name__ == "__main__":
    main()
